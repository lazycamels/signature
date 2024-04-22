import os
import re
from docx import Document

doc = Document('/Users/zhangyidong/Desktop/docx/企业信息.docx')
old_file_path = '/Users/zhangyidong/desktop/docx/'
new_file_path = '/Users/zhangyidong/desktop/new_docx/'

rules =[
    (r'统一社会信用代码：(.*?)企业名称：','-信用代码'),
    (r'企业名称：(.*?)注册号：','-公司名称'),
    (r'法定代表人：(.*?)类型：','-法定代表人'),
    (r'成立日期：(.*?)注册资本：','-成立日期'),
    (r'注册资本：(.*?)核准日期：','-注册资本'),
    (r'住所：(.*?)经营范围：','-公司住所'),
    (r'经营范围：(.*?)提示：','-经营范围')
]

replace_dict = {}
for paragraph in doc.paragraphs:
    text = paragraph.text

    replace_dict1 = {}
    for pattern, variable_name in rules:
        match = re.search(pattern, text, re.DOTALL)
        if match:
            replace_dict1[variable_name] = match.group(1).strip()

        if replace_dict1:
            replace_dict.update(replace_dict1)

def check_and_change(document, replace_dict):
    for para in document.paragraphs:
        for i in range(len(para.runs)):
            for key,value in replace_dict1.items():
                if key in para.runs[i].text:
                    #print(key+'-->'+value)
                    para.runs[i].text = para.runs[i].text.replace(key, value)
    return document

replace_dict1['-报告年度'] = '2023年'
replace_dict1['-报告号'] = '0412号'
replace_dict1['-报告日期'] = '2024年04月17日'

#print(replace_dict)
#print(replace_dict1)

for name in os.listdir(old_file_path):
    print(name)
    old_file = old_file_path+name
    new_file = new_file_path+name
    if old_file.split('.')[1] == 'docx':
        document = Document(old_file)
        document = check_and_change(document, replace_dict1)
        document.save(new_file)
    #print('-'*30)