import os
import fitz
import tkinter as tk
import shutil
from tkinter import filedialog
from docx import Document
from openpyxl import load_workbook

script_dir = os.path.abspath(os.path.dirname(__file__))# 获取当前脚本的绝对路径
cache_dir = os.path.join(script_dir, 'cache')# 在脚本目录下创建一个名为 'cache' 的子目录用于存放缓存文件
if not os.path.exists(cache_dir):
    os.makedirs(cache_dir)

window_1 = tk.Tk()
window_1.withdraw()

file_path = filedialog.askopenfilename(
    title="选择报表文件",
    filetypes=(("Excel files", "*.xlsx"), ("all files", "*.*"))
)
if file_path:
    wb = load_workbook(file_path,data_only=True)
    sheet = wb.worksheets[1]
    doc = Document()
    docx=Document()
    data_pairs = []
    exclude_texts = [
        "流动资产合计",
        "固定资产原价",
        "减：累计折旧",
        "非流动资产合计",
        "资产总计",
        "流动负债合计",
        "非流动负债合计",
        "负债合计",
        "负债和所有者权益（或股东权益）合计",
        "所有者权益（或股东权益）合计",
        "负债和所有者权益（或股东权益）总计",
        "二、营业利润（亏损以“-”号填列）",
        "三、利润总额（亏损总额以“-”号填列）",
        "四、净利润（净亏损以“-”号填列）",
        "城市维护建设税",
        "城镇土地使用税",
        "教育费附加",
        "商品维修费",
        "广告费",
        "业务招待费",
        "利息",
        "政府补助"
        ]

    for row in sheet.iter_rows(min_col=1, max_col=3, values_only=True):
        third_column_value = row[2]
        cell_value = row[0] if row[0] is not None else ''
        if isinstance(third_column_value, (int, float)) and not any(exclude_text in cell_value for exclude_text in exclude_texts):
            formatted_value = '{:.2f}'.format(third_column_value)
            data_pairs.append((cell_value, formatted_value))

    for row in sheet.iter_rows(min_col=5, max_col=7, values_only=True):
        third_column_value = row[2]
        cell_value = row[0] if row[0] is not None else ''
        if isinstance(third_column_value, (int, float)) and not any(exclude_text in cell_value for exclude_text in exclude_texts):
            formatted_value = '{:.2f}'.format(third_column_value)
            data_pairs.append((cell_value, formatted_value))

    for pair in data_pairs:
        paragraph = doc.add_paragraph()
        paragraph.add_run(pair[0])
        paragraph = doc.add_paragraph()
        paragraph.add_run('截止到2023年12月31日' + pair[0] + '为' + pair[1] + '元。')

    sheet = wb.worksheets[2]
    data1_pairs = []

    for row in sheet.iter_rows(min_col=1, max_col=4, values_only=True):
        third_column_value = row[3]
        cell_value = row[0] if row[0] is not None else ''
        if isinstance(third_column_value, (int, float)) and not any(exclude_text in cell_value for exclude_text in exclude_texts):
            formatted_value = '{:.2f}'.format(third_column_value)
            data1_pairs.append((cell_value, formatted_value))

    data_pairs.extend(data1_pairs)

    for pair1 in data1_pairs:
        paragraph1 = doc.add_paragraph()
        paragraph1.add_run(pair1[0])
        paragraph1 = doc.add_paragraph()
        paragraph1.add_run('2023年累计' + pair1[0] + '为' + pair1[1] + '元。')

    doc.save(os.path.join(cache_dir, '明细.docx'))

doc = Document(os.path.join(cache_dir, '明细.docx'))

characters_to_remove = [
    '一、',
    '减：',
    '加：',
    '减：',
    '(或股本）',
    '账面价值'
]
for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        text = run.text
        for char in characters_to_remove:
            text = text.replace(char, '')
        run.text = text

doc.save(os.path.join(cache_dir, '明细.docx'))

docx1 = Document(os.path.join(cache_dir, '明细.docx'))

window_2 = tk.Tk()
window_2.withdraw()

docx2_path = filedialog.askopenfilename(
    title="选择报告文件",
    filetypes=(("Word documents", "*.docx"), ("all files", "*.*"))
)

docx3 = Document(docx2_path)
home_directory = os.path.expanduser('~') #以下三行的作用是获取到设备的桌面路径，并把最终生成的文件保存到桌面
desktop_directory = os.path.join(home_directory, 'Desktop')
output_pdf_path = os.path.join(desktop_directory, '报告附注+明细.docx')

if docx3:
    content_to_insert = "\n".join([para.text for para in docx1.paragraphs])
    for paragraph in docx3.paragraphs:
        if "会计报表主要项目注释" in paragraph.text:
            run = paragraph.add_run('\r')
            run = paragraph.add_run()
            run.text = content_to_insert
            break

docx3.save(output_pdf_path)

os.remove(os.path.join(cache_dir, '明细.docx'))