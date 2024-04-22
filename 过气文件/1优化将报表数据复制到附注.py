# zhangyidong
# 时间：2024/4/14 15:59

import os
import fitz
import tkinter as tk
from tkinter import filedialog
import openpyxl
from docx import Document
from openpyxl import load_workbook

window1 = tk.Tk()
window1.withdraw()

file_path = filedialog.askopenfilename(
    title="选择报表",
    filetypes=(("Excel files", "*.xlsx"), ("all files", "*.*"))
)
if file_path:
    wb = load_workbook(file_path,data_only=True)
    sheet = wb.worksheets[1]
    doc = Document()
    data_pairs = []
    exclude_texts = [
        "流动资产合计",
        "固定资产原价",
        "减：累计折旧",
        "非流动资产合计",
        "资产总计",
        "流动负债合计",
        "非流动负债合计",
        "负债合计",
        "负债和所有者权益（或股东权益）合计",
        "所有者权益（或股东权益）合计",
        "负债和所有者权益（或股东权益）总计",
        "二、营业利润（亏损以“-”号填列）",
        "三、利润总额（亏损总额以“-”号填列）",
        "四、净利润（净亏损以“-”号填列）",
        "城市维护建设税",
        "城镇土地使用税",
        "教育费附加",
        "商品维修费",
        "广告费",
        "业务招待费",
        "利息",
        "政府补助"
        ]

    for row in sheet.iter_rows(min_col=1, max_col=3, values_only=True):
        third_column_value = row[2]
        cell_value = row[0] if row[0] is not None else ''
        if isinstance(third_column_value, (int, float)) and not any(exclude_text in cell_value for exclude_text in exclude_texts):
            formatted_value = '{:.2f}'.format(third_column_value)
            data_pairs.append((cell_value, formatted_value))

    for row in sheet.iter_rows(min_col=5, max_col=7, values_only=True):
        third_column_value = row[2]
        cell_value = row[0] if row[0] is not None else ''
        if isinstance(third_column_value, (int, float)) and not any(exclude_text in cell_value for exclude_text in exclude_texts):
            formatted_value = '{:.2f}'.format(third_column_value)
            data_pairs.append((cell_value, formatted_value))

    for pair in data_pairs:4
        paragraph = doc.add_paragraph()
        paragraph.add_run(pair[0]+'\r' + '截止到2023年12月31日' + pair[0] + '为' + pair[1] + '元。\r')

    sheet = wb.worksheets[2]
    data1_pairs = []

    for row in sheet.iter_rows(min_col=1, max_col=4, values_only=True):
        third_column_value = row[3]
        cell_value = row[0] if row[0] is not None else ''
        if isinstance(third_column_value, (int, float)) and not any(exclude_text in cell_value for exclude_text in exclude_texts):
            formatted_value = '{:.2f}'.format(third_column_value)
            data1_pairs.append((cell_value, formatted_value))

    data_pairs.extend(data1_pairs)

    for pair1 in data1_pairs:
        paragraph1 = doc.add_paragraph()
        paragraph1.add_run(pair1[0] +'\r' + '2023年累计' + pair1[0] + '为' + pair1[1] + '元。\r')

    doc.save('/Users/zhangyidong/Desktop/缓存/附注明细.docx')

doc = Document('/Users/zhangyidong/Desktop/缓存/明细.docx')

copied_text = ""

for paragraph in doc.paragraphs:
    copied_text += paragraph.text + " "

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            copied_text += cell.text + " "

target_doc = Document('/Users/zhangyidong/Desktop/缓存/报告附注.docx')

target_text = "九、或有事项"

target_paragraph = None
for paragraph in target_doc.paragraphs:
    if target_text in paragraph.text:
         target_paragraph = paragraph
         break

if target_paragraph:
     new_paragraph = target_paragraph.insert_paragraph_before(copied_text)

target_doc.save('/Users/zhangyidong/Desktop/缓存/报告附注1.docx')