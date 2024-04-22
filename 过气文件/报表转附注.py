# zhangyidong
# 时间：2024/4/9 12:28
#将表格模板里的特定数字和文字输出到文档模板中

import os
import fitz
import tkinter as tk
from tkinter import filedialog
from docx import Document
from openpyxl import load_workbook

window1 = tk.Tk()
window1.withdraw()

file_path = filedialog.askopenfilename(
    title="选择报表文件",
    filetypes=(("Excel files", "*.xlsx"), ("all files", "*.*"))
)
if file_path:
    wb = load_workbook(file_path,data_only=True)
    sheet = wb.worksheets[1]
    doc = Document()
    data_pairs = []
    exclude_texts = [
        "流动资产合计",
        "固定资产原价",
        "减：累计折旧",
        "非流动资产合计",
        "资产总计",
        "流动负债合计",
        "非流动负债合计",
        "负债合计",
        "负债和所有者权益（或股东权益）合计",
        "所有者权益（或股东权益）合计",
        "负债和所有者权益（或股东权益）总计",
        "二、营业利润（亏损以“-”号填列）",
        "三、利润总额（亏损总额以“-”号填列）",
        "四、净利润（净亏损以“-”号填列）",
        "城市维护建设税",
        "城镇土地使用税",
        "教育费附加",
        "商品维修费",
        "广告费",
        "业务招待费",
        "利息",
        "政府补助"
        ]

    for row in sheet.iter_rows(min_col=1, max_col=3, values_only=True):
        third_column_value = row[2]
        cell_value = row[0] if row[0] is not None else ''
        if isinstance(third_column_value, (int, float)) and not any(exclude_text in cell_value for exclude_text in exclude_texts):
            formatted_value = '{:.2f}'.format(third_column_value)
            data_pairs.append((cell_value, formatted_value))

    for row in sheet.iter_rows(min_col=5, max_col=7, values_only=True):
        third_column_value = row[2]
        cell_value = row[0] if row[0] is not None else ''
        if isinstance(third_column_value, (int, float)) and not any(exclude_text in cell_value for exclude_text in exclude_texts):
            formatted_value = '{:.2f}'.format(third_column_value)
            data_pairs.append((cell_value, formatted_value))

    for pair in data_pairs:
        paragraph = doc.add_paragraph()
        paragraph.add_run(pair[0])
        paragraph = doc.add_paragraph()
        paragraph.add_run('截止到2023年12月31日' + pair[0] + '为' + pair[1] + '元。')

    sheet = wb.worksheets[2]
    data1_pairs = []

    for row in sheet.iter_rows(min_col=1, max_col=4, values_only=True):
        third_column_value = row[3]
        cell_value = row[0] if row[0] is not None else ''
        if isinstance(third_column_value, (int, float)) and not any(exclude_text in cell_value for exclude_text in exclude_texts):
            formatted_value = '{:.2f}'.format(third_column_value)
            data1_pairs.append((cell_value, formatted_value))

    data_pairs.extend(data1_pairs)

    for pair1 in data1_pairs:
        paragraph1 = doc.add_paragraph()
        paragraph1.add_run(pair1[0])
        paragraph1 = doc.add_paragraph()
        paragraph1.add_run('2023年累计' + pair1[0] + '为' + pair1[1] + '元。')

    doc.save('/Users/zhangyidong/Desktop/缓存/明细.docx')

doc = Document('/Users/zhangyidong/Desktop/缓存/明细.docx')

characters_to_remove =[
    '一、',
    '减：',
    '加：',
    '减：',
    '(或股本）',
    '账面价值'
]

# 遍历文档中的所有段落
for paragraph in doc.paragraphs:
    # 遍历段落中的所有运行（run）
    for run in paragraph.runs:
        # 获取运行中的文本
        text = run.text
        # 删除指定的字符
        for char in characters_to_remove:
            text = text.replace(char, '')
        # 更新运行中的文本（这将删除所有指定的字符）
        run.text = text

# 保存修改后的文档
doc.save('/Users/zhangyidong/Desktop/缓存/明细.docx')

# 加载源文档(docx1)和目标文档(docx2)
docx1 = Document('/Users/zhangyidong/Desktop/缓存/明细.docx')
window2 = tk.Tk()
window2.withdraw()

docx2_path = filedialog.askopenfilename(
    title="选择报告文件",
    filetypes=(("Word documents", "*.docx"), ("all files", "*.*"))
)

docx2 = Document(docx2_path)

if docx2:
    # 获取源文档(docx1)中的所有段落文本内容
    content_to_insert = "\n".join([para.text for para in docx1.paragraphs])

    # 搜索关键词“会计报表主要项目注释”
    # 遍历目标文档(docx2)中的所有段落
    for paragraph in docx2.paragraphs:
        # 检查段落文本是否包含关键词
        if "会计报表主要项目注释" in paragraph.text:
            # 找到关键词后，在其后面插入内容
            run = paragraph.add_run('\r')
            run = paragraph.add_run()
            run.text = content_to_insert
            # 不再继续搜索，因为我们假设关键词只会出现一次
            break

    # 保存修改后的目标文档(docx2)
docx2.save('/Users/zhangyidong/Desktop/报告附注+明细.docx')

os.remove('/Users/zhangyidong/Desktop/缓存/明细.docx')